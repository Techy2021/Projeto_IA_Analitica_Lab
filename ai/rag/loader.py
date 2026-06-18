import re
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

from app.config import KNOWLEDGE_BASE_DIR


EXTENSOES_SUPORTADAS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class DocumentoTexto:
    nome_arquivo: str
    caminho: str
    texto: str


class _HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.partes: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.partes.append(data)

    def texto(self) -> str:
        return " ".join(self.partes)


def salvar_upload(uploaded_file, destino: Path = KNOWLEDGE_BASE_DIR) -> Path:
    destino.mkdir(parents=True, exist_ok=True)
    nome_seguro = Path(uploaded_file.name).name
    caminho = destino / nome_seguro
    caminho.write_bytes(uploaded_file.getbuffer())
    return caminho


def carregar_documentos(pasta: Path = KNOWLEDGE_BASE_DIR) -> list[DocumentoTexto]:
    pasta.mkdir(parents=True, exist_ok=True)
    documentos = []
    for caminho in sorted(pasta.iterdir()):
        if caminho.is_file() and caminho.suffix.lower() in EXTENSOES_SUPORTADAS:
            texto = extrair_texto(caminho)
            if texto.strip():
                documentos.append(
                    DocumentoTexto(
                        nome_arquivo=caminho.name,
                        caminho=str(caminho),
                        texto=texto,
                    )
                )
    return documentos


def extrair_texto(caminho: str | Path) -> str:
    caminho = Path(caminho)
    extensao = caminho.suffix.lower()

    if extensao == ".pdf":
        return _extrair_pdf(caminho)
    if extensao == ".docx":
        return _extrair_docx(caminho)
    if extensao == ".txt":
        return caminho.read_text(encoding="utf-8", errors="ignore")
    if extensao == ".md":
        return _extrair_markdown(caminho)

    raise ValueError(
        f"Extensao nao suportada: {extensao}. Use PDF, DOCX, TXT ou MD."
    )


def quebrar_em_chunks(
    texto: str,
    tamanho_chunk: int = 450,
    sobreposicao: int = 70,
) -> list[str]:
    texto_limpo = _normalizar_texto(texto)
    if not texto_limpo:
        return []

    palavras = texto_limpo.split()
    chunks = []
    inicio = 0
    while inicio < len(palavras):
        fim = min(inicio + tamanho_chunk, len(palavras))
        chunk = " ".join(palavras[inicio:fim]).strip()
        if chunk:
            chunks.append(chunk)
        if fim >= len(palavras):
            break
        inicio = max(fim - sobreposicao, inicio + 1)
    return chunks


def _extrair_pdf(caminho: Path) -> str:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as erro:
        raise RuntimeError(
            "Biblioteca pypdf nao instalada. Execute: python -m pip install pypdf"
        ) from erro

    leitor = PdfReader(str(caminho))
    paginas = []
    for pagina in leitor.pages:
        paginas.append(pagina.extract_text() or "")
    return "\n".join(paginas)


def _extrair_docx(caminho: Path) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as erro:
        raise RuntimeError(
            "Biblioteca python-docx nao instalada. Execute: python -m pip install python-docx"
        ) from erro

    documento = Document(str(caminho))
    partes = [paragrafo.text for paragrafo in documento.paragraphs]
    for tabela in documento.tables:
        for linha in tabela.rows:
            partes.append(" | ".join(celula.text for celula in linha.cells))
    return "\n".join(partes)


def _extrair_markdown(caminho: Path) -> str:
    texto_md = caminho.read_text(encoding="utf-8", errors="ignore")
    try:
        import markdown
    except ModuleNotFoundError:
        return re.sub(r"[#*_>`\[\]()-]+", " ", texto_md)

    parser = _HTMLTextExtractor()
    parser.feed(markdown.markdown(texto_md))
    return parser.texto()


def _normalizar_texto(texto: str) -> str:
    texto = texto.replace("\x00", " ")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()
